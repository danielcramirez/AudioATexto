import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
import os
import sys
import wave
import json
import re
import shutil
import subprocess
from datetime import datetime
import urllib.request
import urllib.error
import noisereduce as nr
import numpy as np
import soundfile as sf
from vosk import Model, KaldiRecognizer

# ----------------------------
# Cargar modelo de Vosk
# ----------------------------
# Detectar si estamos ejecutando desde PyInstaller
if getattr(sys, 'frozen', False):
    # Ejecutando desde el .exe
    base_path = sys._MEIPASS
else:
    # Ejecutando desde Python normal
    base_path = os.path.dirname(os.path.abspath(__file__))

MODEL_PATH = os.path.join(base_path, "model")
if not os.path.exists(MODEL_PATH):
    raise FileNotFoundError(f"No se encuentra la carpeta 'model' con el modelo Vosk en: {MODEL_PATH}")

model = Model(MODEL_PATH)

GOOGLE_MODEL = "gemini-2.0-flash"


def llamar_google_api(prompt, api_key, max_output_tokens=1200, response_mime_type=None):
    if not api_key or not api_key.strip():
        raise ValueError("API key de Google no proporcionada.")

    generation_config = {
        "temperature": 0.2,
        "maxOutputTokens": max_output_tokens
    }
    if response_mime_type:
        generation_config["responseMimeType"] = response_mime_type

    payload = {
        "contents": [
            {
                "parts": [{"text": prompt}]
            }
        ],
        "generationConfig": generation_config
    }

    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/{GOOGLE_MODEL}:generateContent"
        f"?key={api_key.strip()}"
    )

    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST"
    )

    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detalle = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Error HTTP de Google API: {e.code} - {detalle}") from e
    except urllib.error.URLError as e:
        raise RuntimeError(f"No se pudo conectar con Google API: {e}") from e

    try:
        partes = data["candidates"][0]["content"]["parts"]
        return "\n".join(p.get("text", "") for p in partes).strip()
    except Exception as e:
        raise RuntimeError(f"Respuesta inesperada de Google API: {data}") from e


# ------------------------------------
# Función para reducir ruido del audio
# ------------------------------------
def mejorar_audio(audio_path):
    data, rate = sf.read(audio_path)
    reduced_noise = nr.reduce_noise(y=data, sr=rate)
    base, ext = os.path.splitext(audio_path)
    out_path = f"{base}_clean.wav"
    sf.write(out_path, reduced_noise, rate)
    return out_path


def preparar_audio_wav(audio_path, text_box=None):
    if audio_path.lower().endswith(".wav"):
        return audio_path, False

    ffmpeg_bin = os.path.join(base_path, "ffmpeg", "bin")
    ffmpeg_exe = os.path.join(ffmpeg_bin, "ffmpeg.exe")
    ffprobe_exe = os.path.join(ffmpeg_bin, "ffprobe.exe")
    if os.path.exists(ffmpeg_exe) and os.path.exists(ffprobe_exe):
        current_path = os.environ.get("PATH", "")
        if ffmpeg_bin.lower() not in current_path.lower():
            os.environ["PATH"] = ffmpeg_bin + os.pathsep + current_path

    try:
        from pydub import AudioSegment
        if os.path.exists(ffmpeg_exe) and os.path.exists(ffprobe_exe):
            AudioSegment.converter = ffmpeg_exe
            AudioSegment.ffmpeg = ffmpeg_exe
            AudioSegment.ffprobe = ffprobe_exe
    except ModuleNotFoundError as e:
        raise RuntimeError(
            f"Dependencia faltante para convertir audio ({e}). Usa Python 3.12 para mp3/m4a."
        ) from e

    try:
        if text_box is not None:
            text_box.insert(tk.END, "Convirtiendo audio a WAV...\n")
            text_box.update()

        audio = AudioSegment.from_file(audio_path)
        base, _ = os.path.splitext(audio_path)
        temp_wav = f"{base}_temp.wav"
        audio.export(temp_wav, format="wav")
        return temp_wav, True
    except Exception as e:
        err = str(e).lower()
        if "ffmpeg" in err or "ffprobe" in err:
            raise RuntimeError("No se pudo convertir el audio. Instala FFmpeg y agrega ffmpeg/bin al PATH.") from e
        raise RuntimeError(f"Error convirtiendo audio: {e}") from e


# ------------------------------------
# Verificar calidad del audio
# ------------------------------------
def verificar_calidad(audio_path):
    try:
        data, rate = sf.read(audio_path)
        rms = np.sqrt(np.mean(data ** 2))
        if rms < 0.01:
            return "Calidad baja: volumen muy bajo."
        elif rms > 0.5:
            return "Calidad regular: posible saturación."
        else:
            return "Calidad buena: señal adecuada."
    except Exception as e:
        return f"Error analizando audio: {e}"


# ------------------------------------
# Diarización (identificación de voces)
# ------------------------------------
def identificar_voces(audio_path):
    try:
        # pyAudioAnalysis depende de 'aifc', eliminado en Python 3.13.
        # Se importa aqui para no bloquear toda la app al iniciar.
        from pyAudioAnalysis import audioSegmentation as aS
        diarization_result = aS.speaker_diarization(audio_path, 2)
        if isinstance(diarization_result, tuple) and len(diarization_result) >= 2:
            return diarization_result[1]
        return diarization_result
    except ModuleNotFoundError as e:
        if getattr(e, "name", "") == "aifc":
            return "Diarizacion no disponible en Python 3.13 por dependencia 'aifc'. Usa Python 3.12 o anterior para esta funcion."
        return f"No se pudo hacer diarizacion: falta dependencia ({e})."
    except Exception:
        return "No se pudo hacer diarización, audio muy corto o ruidoso."


# ------------------------------------
# Transcribir audio
# ------------------------------------
def transcribir(audio_path, text_box):
    temp_created = False
    try:
        audio_path, temp_created = preparar_audio_wav(audio_path, text_box)
    except RuntimeError as e:
        text_box.insert(tk.END, f"{e}\n")
        return ""

    try:
        wf = wave.open(audio_path, "rb")
    except Exception as e:
        text_box.insert(tk.END, f"Error abriendo archivo de audio: {e}\n")
        return ""
    rec = KaldiRecognizer(model, wf.getframerate())
    rec.SetWords(True)

    text_box.insert(tk.END, "Iniciando transcripción...\n")
    text_box.update()

    result_text = ""

    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break

        if rec.AcceptWaveform(data):
            result = json.loads(rec.Result())
            if "text" in result:
                line = result["text"] + "\n"
                text_box.insert(tk.END, line)
                text_box.update()
                result_text += line

    final = json.loads(rec.FinalResult())
    if "text" in final:
        text_box.insert(tk.END, "\nFIN DE TRANSCRIPCIÓN\n")
        result_text += final["text"]

    if temp_created and os.path.exists(audio_path):
        try:
            os.remove(audio_path)
        except Exception:
            pass

    return result_text


def resumir_texto_con_google(texto, api_key):
    if not texto or not texto.strip():
        raise ValueError("No hay texto para resumir.")

    texto_limpio = texto.strip()
    if len(texto_limpio) > 12000:
        texto_limpio = texto_limpio[:12000]

    prompt = (
        "Eres un asistente que resume reuniones en español. "
        "Devuelve una respuesta clara con estas secciones:\n"
        "1) Resumen ejecutivo\n"
        "2) Decisiones tomadas\n"
        "3) Tareas y responsables\n"
        "4) Riesgos o pendientes\n"
        "5) Proximos pasos\n\n"
        "Texto de la reunion:\n"
        f"{texto_limpio}"
    )

    texto_resumen = llamar_google_api(prompt, api_key, max_output_tokens=1200)

    if not texto_resumen:
        raise RuntimeError("Google API no devolvio texto de resumen.")

    return texto_resumen


def _extraer_json_de_texto(texto):
    if not texto:
        raise ValueError("No se recibio contenido JSON.")

    limpio = texto.strip()
    if limpio.startswith("```"):
        m = re.search(r"```(?:json)?\s*(.*?)```", limpio, flags=re.DOTALL | re.IGNORECASE)
        if m:
            limpio = m.group(1).strip()

    try:
        return json.loads(limpio)
    except Exception:
        pass

    ini = limpio.find("{")
    fin = limpio.rfind("}")
    if ini >= 0 and fin > ini:
        candidato = limpio[ini:fin + 1]
        return json.loads(candidato)

    raise ValueError("No fue posible parsear JSON de la respuesta del modelo.")


def _valor_texto(v, default="N/A"):
    if v is None:
        return default
    s = str(v).strip()
    return s if s else default


def _latex_escape(texto):
    s = _valor_texto(texto, "")
    reemplazos = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for k, v in reemplazos.items():
        s = s.replace(k, v)
    return s


def generar_datos_acta_local(transcripcion):
    hoy = datetime.now().strftime("%d/%m/%Y")
    resumen_local = resumir_texto_local(transcripcion)
    desarrollo = [x.strip("- ").strip() for x in resumen_local.splitlines() if x.strip().startswith("-")]
    if not desarrollo:
        desarrollo = ["Se realizo la reunion y se revisaron los temas previstos."]

    return {
        "acta_numero": "N/A",
        "fecha": hoy,
        "objetivo": "Revisar temas de la reunion y dejar compromisos.",
        "entidad_organizadora": "N/A",
        "lugar": "N/A",
        "lider_reunion": "N/A",
        "hora_programada_desde": "00:00",
        "hora_programada_hasta": "00:00",
        "hora_inicio": "00:00",
        "hora_terminacion": "00:00",
        "participantes": [],
        "orden_dia": [
            "Revision de compromisos pendientes.",
            "Desarrollo de temas principales.",
            "Definicion de compromisos y proximos pasos."
        ],
        "revision_compromisos": [
            {
                "actividad": "N/A",
                "responsable": "N/A",
                "cumple": "NO",
                "nueva_fecha": "Seleccion",
                "observaciones": ""
            }
        ],
        "desarrollo_temas": desarrollo,
        "observaciones_conclusiones": [
            "Anexar Listado de Asistencia a Reuniones (Formato GEDO-FR06) firmada por los participantes."
        ],
        "compromisos": [
            {
                "actividad": "N/A",
                "responsable": "N/A",
                "fecha_limite": "Seleccion",
                "observaciones": ""
            }
        ],
        "firmas": [
            {
                "nombre": "ANDRES FERNANDO AGUDELO AGUILAR",
                "cargo": "Director Administrativo y Financiero"
            },
            {
                "nombre": "LUISA FERNANDA GONZALEZ MOZO",
                "cargo": "Jefe Oficina Asesora de Planeacion y Control de Riesgos"
            }
        ],
        "control_cambios": [
            {
                "version": "1",
                "fecha": "14 de septiembre de 2020",
                "descripcion": "Version Inicial",
                "asesor": "Andrea Catalina Cuesta Ruiz"
            }
        ]
    }


def generar_datos_acta_con_google(transcripcion, api_key):
    if not transcripcion or not transcripcion.strip():
        raise ValueError("No hay transcripcion para generar el acta.")

    texto_limpio = transcripcion.strip()
    if len(texto_limpio) > 15000:
        texto_limpio = texto_limpio[:15000]

    prompt = (
        "Eres un asistente experto en redaccion de actas en Colombia. "
        "Debes analizar la transcripcion y devolver SOLO JSON valido, sin markdown, sin texto adicional. "
        "Usa datos inferidos del texto y cuando falte informacion usa 'N/A'.\n\n"
        "Estructura exacta requerida:\n"
        "{\n"
        "  \"acta_numero\": \"str\",\n"
        "  \"fecha\": \"dd/mm/aaaa o N/A\",\n"
        "  \"objetivo\": \"str\",\n"
        "  \"entidad_organizadora\": \"str\",\n"
        "  \"lugar\": \"str\",\n"
        "  \"lider_reunion\": \"str\",\n"
        "  \"hora_programada_desde\": \"hh:mm\",\n"
        "  \"hora_programada_hasta\": \"hh:mm\",\n"
        "  \"hora_inicio\": \"hh:mm\",\n"
        "  \"hora_terminacion\": \"hh:mm\",\n"
        "  \"participantes\": [\n"
        "    {\"nombre\":\"str\",\"cargo\":\"str\",\"dependencia\":\"str\",\"asiste\":\"SI|NO\",\"tipo_participante\":\"str\"}\n"
        "  ],\n"
        "  \"orden_dia\": [\"str\"],\n"
        "  \"revision_compromisos\": [\n"
        "    {\"actividad\":\"str\",\"responsable\":\"str\",\"cumple\":\"SI|NO\",\"nueva_fecha\":\"str\",\"observaciones\":\"str\"}\n"
        "  ],\n"
        "  \"desarrollo_temas\": [\"str\"],\n"
        "  \"observaciones_conclusiones\": [\"str\"],\n"
        "  \"compromisos\": [\n"
        "    {\"actividad\":\"str\",\"responsable\":\"str\",\"fecha_limite\":\"str\",\"observaciones\":\"str\"}\n"
        "  ],\n"
        "  \"firmas\": [\n"
        "    {\"nombre\":\"str\",\"cargo\":\"str\"},\n"
        "    {\"nombre\":\"str\",\"cargo\":\"str\"}\n"
        "  ],\n"
        "  \"control_cambios\": [\n"
        "    {\"version\":\"str\",\"fecha\":\"str\",\"descripcion\":\"str\",\"asesor\":\"str\"}\n"
        "  ]\n"
        "}\n\n"
        "Transcripcion de reunion:\n"
        f"{texto_limpio}"
    )

    respuesta = llamar_google_api(
        prompt,
        api_key,
        max_output_tokens=2600,
        response_mime_type="application/json"
    )
    return _extraer_json_de_texto(respuesta)


def generar_acta_latex(datos):
    acta_numero = _latex_escape(_valor_texto(datos.get("acta_numero")))
    fecha = _latex_escape(_valor_texto(datos.get("fecha")))
    objetivo = _latex_escape(_valor_texto(datos.get("objetivo")))
    entidad = _latex_escape(_valor_texto(datos.get("entidad_organizadora")))
    lugar = _latex_escape(_valor_texto(datos.get("lugar")))
    lider = _latex_escape(_valor_texto(datos.get("lider_reunion")))
    hp_desde = _latex_escape(_valor_texto(datos.get("hora_programada_desde"), "00:00"))
    hp_hasta = _latex_escape(_valor_texto(datos.get("hora_programada_hasta"), "00:00"))
    h_inicio = _latex_escape(_valor_texto(datos.get("hora_inicio"), "00:00"))
    h_fin = _latex_escape(_valor_texto(datos.get("hora_terminacion"), "00:00"))

    participantes = datos.get("participantes") or []
    participantes_rows = []
    for i in range(8):
        p = participantes[i] if i < len(participantes) and isinstance(participantes[i], dict) else {}
        asiste = _valor_texto(p.get("asiste"), "NO").upper()
        asiste_cell = "$\\blacksquare$ \\hspace{0.3cm} $\\square$" if asiste == "SI" else "$\\square$ \\hspace{0.3cm} $\\blacksquare$"
        participantes_rows.append(
            f"{i + 1} & {_latex_escape(_valor_texto(p.get('nombre'), ''))} & {_latex_escape(_valor_texto(p.get('cargo'), ''))} & "
            f"{_latex_escape(_valor_texto(p.get('dependencia'), ''))} & {asiste_cell} & {_latex_escape(_valor_texto(p.get('tipo_participante'), 'N/A'))} \\\\ \\hline"
        )

    orden_dia = datos.get("orden_dia") or []
    if not orden_dia:
        orden_dia = [
            "Revision del mecanismo para la implementacion de formatos contractuales y de pago."
        ]
    orden_rows = [f"{idx}. {_latex_escape(_valor_texto(item))} \\\\ \\hline" for idx, item in enumerate(orden_dia, 1)]

    rev_comp = datos.get("revision_compromisos") or []
    if not rev_comp:
        rev_comp = [{"actividad": "N/A", "responsable": "", "cumple": "NO", "nueva_fecha": "Seleccion", "observaciones": ""}]
    rev_rows = []
    for idx, item in enumerate(rev_comp, 1):
        if not isinstance(item, dict):
            item = {}
        cumple = _valor_texto(item.get("cumple"), "NO").upper()
        si_mark = "$\\blacksquare$" if cumple == "SI" else "$\\square$"
        no_mark = "$\\blacksquare$" if cumple != "SI" else "$\\square$"
        rev_rows.append(
            f"{idx} & {_latex_escape(_valor_texto(item.get('actividad')))} & {_latex_escape(_valor_texto(item.get('responsable'), ''))} & "
            f"{si_mark} & {no_mark} & {_latex_escape(_valor_texto(item.get('nueva_fecha'), 'Seleccion'))} & "
            f"{_latex_escape(_valor_texto(item.get('observaciones'), ''))} \\\\ \\hline"
        )

    desarrollo = datos.get("desarrollo_temas") or ["N/A"]
    desarrollo_rows = [f"{_latex_escape(_valor_texto(x))} \\\\ \\hline" for x in desarrollo]

    observaciones = datos.get("observaciones_conclusiones") or [
        "Anexar Listado de Asistencia a Reuniones (Formato GEDO-FR06) firmada por los participantes."
    ]
    observaciones_rows = [f"{_latex_escape(_valor_texto(x))} \\\\ \\hline" for x in observaciones]

    compromisos = datos.get("compromisos") or [{"actividad": "", "responsable": "", "fecha_limite": "Seleccion", "observaciones": ""}]
    comp_rows = []
    for idx, item in enumerate(compromisos, 1):
        if not isinstance(item, dict):
            item = {}
        comp_rows.append(
            f"{idx} & {_latex_escape(_valor_texto(item.get('actividad'), ''))} & {_latex_escape(_valor_texto(item.get('responsable'), ''))} & "
            f"{_latex_escape(_valor_texto(item.get('fecha_limite'), 'Seleccion'))} & {_latex_escape(_valor_texto(item.get('observaciones'), ''))} \\\\ \\hline"
        )

    firmas = datos.get("firmas") or []
    firma1 = firmas[0] if len(firmas) > 0 and isinstance(firmas[0], dict) else {}
    firma2 = firmas[1] if len(firmas) > 1 and isinstance(firmas[1], dict) else {}

    control = datos.get("control_cambios") or [{
        "version": "1",
        "fecha": "14 de septiembre de 2020",
        "descripcion": "Version Inicial",
        "asesor": "Andrea Catalina Cuesta Ruiz"
    }]
    control_rows = []
    for item in control:
        if not isinstance(item, dict):
            item = {}
        control_rows.append(
            f"{_latex_escape(_valor_texto(item.get('version'), '1'))} & {_latex_escape(_valor_texto(item.get('fecha')))} & "
            f"{_latex_escape(_valor_texto(item.get('descripcion')))} & {_latex_escape(_valor_texto(item.get('asesor')))} \\\\ \\hline"
        )

    lines = [
        "\\documentclass[10pt,a4paper]{article}",
        "\\usepackage[utf8]{inputenc}",
        "\\usepackage[spanish]{babel}",
        "\\usepackage[left=1.5cm,right=1.5cm,top=4cm,bottom=2cm,headheight=2.5cm]{geometry}",
        "\\usepackage{tabularx}",
        "\\usepackage{multirow}",
        "\\usepackage{graphicx}",
        "\\usepackage{array}",
        "\\usepackage{fancyhdr}",
        "\\usepackage{amssymb}",
        "\\usepackage{helvet}",
        "\\renewcommand{\\familydefault}{\\sfdefault}",
        "\\usepackage{lastpage}",
        "",
        "\\pagestyle{fancy}",
        "\\fancyhf{}",
        "\\renewcommand{\\headrulewidth}{0pt}",
        "\\fancyhead[C]{",
        "    \\renewcommand{\\arraystretch}{1.4}",
        "    \\begin{tabularx}{\\textwidth}{|c|X|m{4cm}|}",
        "    \\hline",
        "    \\multirow{4}{*}{\\includegraphics[width=2.5cm,keepaspectratio]{logo.png}}",
        "    & \\textbf{PROCESO} \\newline GESTION DE DESARROLLO ORGANIZACIONAL & \\textbf{Codigo:} GEDO-FR05 \\\\ \\cline{2-3}",
        "    & \\multirow{2}{*}{\\textbf{FORMATO} \\newline Acta de Reunion} & \\textbf{Version:} 02 \\\\ \\cline{3-3}",
        "    & & \\textbf{Fecha:} 14/09/2020 \\\\ \\hline",
        "    \\end{tabularx}",
        "}",
        "\\fancyfoot[R]{Pagina \\thepage\\ de \\pageref{LastPage}}",
        "",
        "\\begin{document}",
        "\\renewcommand{\\arraystretch}{1.3}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|X|X|X|}",
        "\\hline",
        f"\\textbf{{Acta N$^o$:}} {acta_numero} & \\textbf{{Fecha:}} {fecha} & \\textbf{{Objetivo:}} {objetivo} \\\\ \\hline",
        f"\\multicolumn{{3}}{{|l|}}{{\\textbf{{Nombre de la Dependencia, Proceso o Entidad que organiza la reunion:}} {entidad}}} \\\\ \\hline",
        f"\\textbf{{Lugar:}} {lugar} & \\multicolumn{{2}}{{l|}}{{\\textbf{{Lider de la reunion:}} {lider}}} \\\\ \\hline",
        f"\\textbf{{Hora Programada:}} \\newline De: {hp_desde} A {hp_hasta} & \\textbf{{Hora de Inicio:}} \\newline {h_inicio} & \\textbf{{Hora de Terminacion:}} \\newline {h_fin} \\\\ \\hline",
        "\\end{tabularx}",
        "",
        "\\vspace{0.4cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|c|X|X|X|c|c|}",
        "\\hline",
        "\\multirow{2}{*}{\\textbf{No}} & \\multirow{2}{*}{\\textbf{Nombre*}} & \\multirow{2}{*}{\\textbf{Cargo}} & \\multirow{2}{*}{\\textbf{Dependencia o entidad}} & \\textbf{Asiste} & \\multirow{2}{*}{\\textbf{Tipo de Participante}} \\\\ \\cline{5-5}",
        "& & & & SI \\hspace{0.2cm} NO & \\\\ \\hline",
        *participantes_rows,
        "\\end{tabularx}",
        "",
        "\\vspace{0.4cm}",
        "",
        "\\noindent\\textbf{ORDEN DEL DIA}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|X|}",
        "\\hline",
        *orden_rows,
        "\\end{tabularx}",
        "",
        "\\vspace{0.4cm}",
        "",
        "\\noindent\\textbf{DESARROLLO DEL ORDEN DEL DIA}",
        "\\vspace{0.2cm}",
        "",
        "\\noindent\\textbf{1. REVISION DE COMPROMISOS PENDIENTES:}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|c|X|X|c|c|c|X|}",
        "\\hline",
        "\\multirow{2}{*}{\\textbf{N$^o$}} & \\multirow{2}{*}{\\textbf{Actividad}} & \\multirow{2}{*}{\\textbf{Responsable}} & \\multicolumn{2}{c|}{\\textbf{Cumple}} & \\multirow{2}{*}{\\textbf{Nueva fecha}} & \\multirow{2}{*}{\\textbf{Observaciones}} \\\\ \\cline{4-5}",
        "& & & \\textbf{SI} & \\textbf{NO} & & \\\\ \\hline",
        *rev_rows,
        "\\end{tabularx}",
        "",
        "\\vspace{0.4cm}",
        "",
        "\\noindent\\textbf{2. DESARROLLO DE LOS TEMAS:}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|X|}",
        "\\hline",
        *desarrollo_rows,
        "\\end{tabularx}",
        "",
        "\\vspace{0.4cm}",
        "",
        "\\noindent\\textbf{3. OBSERVACIONES Y CONCLUSIONES:}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|X|}",
        "\\hline",
        *observaciones_rows,
        "\\end{tabularx}",
        "",
        "\\newpage",
        "",
        "\\noindent\\textbf{4. ESTABLECIMIENTO DE COMPROMISOS:}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|c|X|X|c|X|}",
        "\\hline",
        "\\textbf{N$^o$} & \\textbf{Actividad} & \\textbf{Responsable} & \\textbf{Fecha Limite} & \\textbf{Observaciones} \\\\ \\hline",
        *comp_rows,
        "\\end{tabularx}",
        "",
        "\\vspace{2.5cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{X X}",
        "\\rule{6cm}{0.5pt} & \\rule{6cm}{0.5pt} \\\\",
        f"\\textbf{{{_latex_escape(_valor_texto(firma1.get('nombre'), 'ANDRES FERNANDO AGUDELO AGUILAR'))}}} & \\textbf{{{_latex_escape(_valor_texto(firma2.get('nombre'), 'LUISA FERNANDA GONZALEZ MOZO'))}}} \\\\",
        f"{_latex_escape(_valor_texto(firma1.get('cargo'), 'Director Administrativo y Financiero'))} & {_latex_escape(_valor_texto(firma2.get('cargo'), 'Jefe Oficina Asesora de Planeacion y Control de Riesgos'))} \\\\",
        "\\end{tabularx}",
        "",
        "\\vspace{1.5cm}",
        "",
        "\\noindent\\textbf{5. CONTROL DE CAMBIOS}",
        "\\vspace{0.1cm}",
        "",
        "\\noindent",
        "\\begin{tabularx}{\\textwidth}{|c|m{3cm}|X|X|}",
        "\\hline",
        "\\textbf{Version} & \\textbf{Fecha} & \\textbf{Descripcion del cambio} & \\textbf{Asesor del proceso} \\\\ \\hline",
        *control_rows,
        "\\end{tabularx}",
        "",
        "\\end{document}",
    ]

    return "\n".join(lines)


def _segmentar_oraciones(texto):
    limpio = " ".join((texto or "").split())
    if not limpio:
        return []
    partes = re.split(r"(?<=[\.!\?])\s+", limpio)
    return [p.strip() for p in partes if p and len(p.strip()) > 20]


def resumir_texto_local(texto):
    oraciones = _segmentar_oraciones(texto)
    if not oraciones:
        return "No fue posible generar resumen local: transcripción vacía o demasiado corta."

    texto_minus = " ".join(oraciones).lower()

    claves_decision = ["decid", "acord", "aprob", "se define", "se confirma"]
    claves_tarea = ["tarea", "responsable", "encarg", "debe", "compromiso", "accion"]
    claves_riesgo = ["riesgo", "bloque", "pendiente", "problema", "retras", "falta"]
    claves_siguientes = ["siguiente", "proximo", "continuar", "luego", "plan"]

    decisiones = [o for o in oraciones if any(k in o.lower() for k in claves_decision)]
    tareas = [o for o in oraciones if any(k in o.lower() for k in claves_tarea)]
    riesgos = [o for o in oraciones if any(k in o.lower() for k in claves_riesgo)]
    siguientes = [o for o in oraciones if any(k in o.lower() for k in claves_siguientes)]

    resumen_base = oraciones[:4]
    if len(texto_minus) > 1500 and len(oraciones) > 6:
        resumen_base = oraciones[:6]

    def bullets(items, max_items=5):
        if not items:
            return "- Sin elementos detectados automaticamente."
        return "\n".join(f"- {x}" for x in items[:max_items])

    return (
        "1) Resumen ejecutivo\n"
        f"{bullets(resumen_base, max_items=6)}\n\n"
        "2) Decisiones tomadas\n"
        f"{bullets(decisiones)}\n\n"
        "3) Tareas y responsables\n"
        f"{bullets(tareas)}\n\n"
        "4) Riesgos o pendientes\n"
        f"{bullets(riesgos)}\n\n"
        "5) Proximos pasos\n"
        f"{bullets(siguientes)}\n\n"
        "Nota: resumen generado en modo local por falta de cuota o conectividad de Google API."
    )


# ------------------------------------
# Interfaz gráfica
# ------------------------------------

def seleccionar_audio():
    global audio_file
    audio_file = filedialog.askopenfilename(
        title="Seleccionar archivo de audio",
        filetypes=[
            ("Archivos de Audio", "*.wav *.ogg *.mp3 *.m4a"),
            ("Archivos WAV", "*.wav"),
            ("Archivos OGG", "*.ogg"),
            ("Archivos MP3", "*.mp3"),
            ("Archivos M4A", "*.m4a")
        ]
    )

    if audio_file:
        text_box.insert(tk.END, f"Audio seleccionado: {audio_file}\n\n")


def procesar_transcripcion():
    global ultima_transcripcion
    if not audio_file:
        messagebox.showerror("Error", "Debe seleccionar un archivo de audio.")
        return

    text_box.delete(1.0, tk.END)
    resultado = transcribir(audio_file, text_box)

    if resultado:
        ultima_transcripcion = resultado
        base_name = os.path.splitext(audio_file)[0]
        txt_path = f"{base_name}_transcripcion.txt"
        try:
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(resultado)
            messagebox.showinfo("Éxito", f"Transcripción guardada en:\n{txt_path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el archivo de texto: {e}")


def procesar_resumen_reunion():
    global ultima_transcripcion

    if not ultima_transcripcion or not ultima_transcripcion.strip():
        messagebox.showerror("Error", "Primero debes transcribir un audio para poder resumir.")
        return

    api_key = api_key_var.get().strip() or os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        messagebox.showerror("Error", "Ingresa la API key de Google o configura GEMINI_API_KEY.")
        return

    text_box.insert(tk.END, "\nGenerando resumen de reunión con Google AI...\n")
    text_box.update()

    try:
        resumen = resumir_texto_con_google(ultima_transcripcion, api_key)
        text_box.insert(tk.END, "\n===== RESUMEN DE REUNION =====\n")
        text_box.insert(tk.END, resumen + "\n")
        text_box.update()

        if audio_file:
            base_name = os.path.splitext(audio_file)[0]
            resumen_path = f"{base_name}_resumen_reunion.txt"
            with open(resumen_path, "w", encoding="utf-8") as f:
                f.write(resumen)
            messagebox.showinfo("Éxito", f"Resumen guardado en:\n{resumen_path}")
    except Exception as e:
        err_txt = str(e).lower()
        if "429" in err_txt or "quota" in err_txt or "resource_exhausted" in err_txt:
            resumen = resumir_texto_local(ultima_transcripcion)
            text_box.insert(tk.END, "\n===== RESUMEN DE REUNION (MODO LOCAL) =====\n")
            text_box.insert(tk.END, resumen + "\n")
            text_box.update()

            if audio_file:
                base_name = os.path.splitext(audio_file)[0]
                resumen_path = f"{base_name}_resumen_reunion_local.txt"
                with open(resumen_path, "w", encoding="utf-8") as f:
                    f.write(resumen)
                messagebox.showwarning(
                    "Sin cuota en Google API",
                    "Google API reporto cuota agotada (429). Se genero un resumen local de respaldo en:\n"
                    f"{resumen_path}"
                )
            return

        messagebox.showerror("Error", f"No se pudo generar el resumen: {e}")


def procesar_generar_acta_latex():
    global ultima_transcripcion

    if not ultima_transcripcion or not ultima_transcripcion.strip():
        messagebox.showerror("Error", "Primero debes transcribir un audio para generar el acta.")
        return

    api_key = api_key_var.get().strip() or os.environ.get("GEMINI_API_KEY", "").strip()

    text_box.insert(tk.END, "\nGenerando acta en LaTeX...\n")
    text_box.update()

    try:
        if api_key:
            datos = generar_datos_acta_con_google(ultima_transcripcion, api_key)
        else:
            datos = generar_datos_acta_local(ultima_transcripcion)

        latex_acta = generar_acta_latex(datos)
        output_path = os.path.join(base_path, "acta_reunion_generada.tex")
        if audio_file:
            base_name = os.path.splitext(audio_file)[0]
            output_path = f"{base_name}_acta_reunion.tex"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(latex_acta)

        text_box.insert(tk.END, f"Acta LaTeX generada: {output_path}\n")
        text_box.update()
        messagebox.showinfo("Exito", f"Acta generada en:\n{output_path}")
    except Exception as e:
        err_txt = str(e).lower()
        if "429" in err_txt or "quota" in err_txt or "resource_exhausted" in err_txt:
            try:
                datos = generar_datos_acta_local(ultima_transcripcion)
                latex_acta = generar_acta_latex(datos)
                output_path = os.path.join(base_path, "acta_reunion_generada_local.tex")
                if audio_file:
                    base_name = os.path.splitext(audio_file)[0]
                    output_path = f"{base_name}_acta_reunion_local.tex"
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(latex_acta)
                text_box.insert(tk.END, f"Acta LaTeX local generada: {output_path}\n")
                text_box.update()
                messagebox.showwarning(
                    "Sin cuota en Google API",
                    "Google API reporto cuota agotada (429). Se genero el acta en modo local en:\n"
                    f"{output_path}"
                )
                return
            except Exception as e_local:
                messagebox.showerror("Error", f"No se pudo generar el acta local: {e_local}")
                return

        messagebox.showerror("Error", f"No se pudo generar el acta: {e}")


def _generar_y_guardar_acta(usar_google=True):
    api_key = api_key_var.get().strip() or os.environ.get("GEMINI_API_KEY", "").strip()
    uso_local_por_cuota = False

    if usar_google and api_key:
        try:
            datos = generar_datos_acta_con_google(ultima_transcripcion, api_key)
            output_path = os.path.join(base_path, "acta_reunion_generada.tex")
            if audio_file:
                base_name = os.path.splitext(audio_file)[0]
                output_path = f"{base_name}_acta_reunion.tex"
        except Exception as e:
            err_txt = str(e).lower()
            if "429" in err_txt or "quota" in err_txt or "resource_exhausted" in err_txt:
                uso_local_por_cuota = True
                datos = generar_datos_acta_local(ultima_transcripcion)
                output_path = os.path.join(base_path, "acta_reunion_generada_local.tex")
                if audio_file:
                    base_name = os.path.splitext(audio_file)[0]
                    output_path = f"{base_name}_acta_reunion_local.tex"
            else:
                raise
    else:
        datos = generar_datos_acta_local(ultima_transcripcion)
        output_path = os.path.join(base_path, "acta_reunion_generada_local.tex")
        if audio_file:
            base_name = os.path.splitext(audio_file)[0]
            output_path = f"{base_name}_acta_reunion_local.tex"

    latex_acta = generar_acta_latex(datos)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(latex_acta)

    return output_path, uso_local_por_cuota, datos


def _guardar_acta_docx(datos, docx_path):
    try:
        from docx import Document  # type: ignore[reportMissingImports]
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "Falta dependencia para Word: python-docx. Ejecuta: pip install -r requirements.txt"
        ) from e

    def t(v, default="N/A"):
        return _valor_texto(v, default)

    doc = Document()
    doc.add_heading("ACTA DE REUNION", level=1)

    info = doc.add_table(rows=0, cols=2)
    info.style = "Table Grid"

    generales = [
        ("Acta N°", t(datos.get("acta_numero"))),
        ("Fecha", t(datos.get("fecha"))),
        ("Objetivo", t(datos.get("objetivo"))),
        ("Entidad organizadora", t(datos.get("entidad_organizadora"))),
        ("Lugar", t(datos.get("lugar"))),
        ("Lider de la reunion", t(datos.get("lider_reunion"))),
        ("Hora programada", f"De {t(datos.get('hora_programada_desde'), '00:00')} a {t(datos.get('hora_programada_hasta'), '00:00')}"),
        ("Hora inicio", t(datos.get("hora_inicio"), "00:00")),
        ("Hora terminacion", t(datos.get("hora_terminacion"), "00:00")),
    ]
    for k, v in generales:
        row = info.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_heading("Participantes", level=2)
    participantes = datos.get("participantes") or []
    part_table = doc.add_table(rows=1, cols=6)
    part_table.style = "Table Grid"
    headers = ["No", "Nombre", "Cargo", "Dependencia", "Asiste", "Tipo"]
    for i, h in enumerate(headers):
        part_table.rows[0].cells[i].text = h

    if participantes:
        for idx, p in enumerate(participantes, 1):
            if not isinstance(p, dict):
                p = {}
            r = part_table.add_row().cells
            r[0].text = str(idx)
            r[1].text = t(p.get("nombre"), "")
            r[2].text = t(p.get("cargo"), "")
            r[3].text = t(p.get("dependencia"), "")
            r[4].text = t(p.get("asiste"), "NO")
            r[5].text = t(p.get("tipo_participante"), "N/A")
    else:
        r = part_table.add_row().cells
        r[0].text = "1"
        r[1].text = ""
        r[2].text = ""
        r[3].text = ""
        r[4].text = "NO"
        r[5].text = "N/A"

    def agregar_lista(titulo, items):
        doc.add_heading(titulo, level=2)
        valores = items if isinstance(items, list) and items else ["N/A"]
        for item in valores:
            doc.add_paragraph(t(item), style="List Bullet")

    agregar_lista("Orden del dia", datos.get("orden_dia") or [])

    doc.add_heading("Revision de compromisos pendientes", level=2)
    rev = datos.get("revision_compromisos") or []
    rev_table = doc.add_table(rows=1, cols=6)
    rev_table.style = "Table Grid"
    rev_headers = ["No", "Actividad", "Responsable", "Cumple", "Nueva fecha", "Observaciones"]
    for i, h in enumerate(rev_headers):
        rev_table.rows[0].cells[i].text = h
    if not rev:
        rev = [{"actividad": "N/A", "responsable": "", "cumple": "NO", "nueva_fecha": "Seleccion", "observaciones": ""}]
    for idx, item in enumerate(rev, 1):
        if not isinstance(item, dict):
            item = {}
        r = rev_table.add_row().cells
        r[0].text = str(idx)
        r[1].text = t(item.get("actividad"))
        r[2].text = t(item.get("responsable"), "")
        r[3].text = t(item.get("cumple"), "NO")
        r[4].text = t(item.get("nueva_fecha"), "Seleccion")
        r[5].text = t(item.get("observaciones"), "")

    agregar_lista("Desarrollo de los temas", datos.get("desarrollo_temas") or [])
    agregar_lista("Observaciones y conclusiones", datos.get("observaciones_conclusiones") or [])

    doc.add_heading("Establecimiento de compromisos", level=2)
    comp = datos.get("compromisos") or []
    comp_table = doc.add_table(rows=1, cols=5)
    comp_table.style = "Table Grid"
    comp_headers = ["No", "Actividad", "Responsable", "Fecha limite", "Observaciones"]
    for i, h in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = h
    if not comp:
        comp = [{"actividad": "", "responsable": "", "fecha_limite": "Seleccion", "observaciones": ""}]
    for idx, item in enumerate(comp, 1):
        if not isinstance(item, dict):
            item = {}
        r = comp_table.add_row().cells
        r[0].text = str(idx)
        r[1].text = t(item.get("actividad"), "")
        r[2].text = t(item.get("responsable"), "")
        r[3].text = t(item.get("fecha_limite"), "Seleccion")
        r[4].text = t(item.get("observaciones"), "")

    doc.add_heading("Firmas", level=2)
    firmas = datos.get("firmas") or []
    f1 = firmas[0] if len(firmas) > 0 and isinstance(firmas[0], dict) else {}
    f2 = firmas[1] if len(firmas) > 1 and isinstance(firmas[1], dict) else {}
    doc.add_paragraph(f"1) {t(f1.get('nombre'), 'ANDRES FERNANDO AGUDELO AGUILAR')} - {t(f1.get('cargo'), 'Director Administrativo y Financiero')}")
    doc.add_paragraph(f"2) {t(f2.get('nombre'), 'LUISA FERNANDA GONZALEZ MOZO')} - {t(f2.get('cargo'), 'Jefe Oficina Asesora de Planeacion y Control de Riesgos')}")

    doc.add_heading("Control de cambios", level=2)
    control = datos.get("control_cambios") or [{"version": "1", "fecha": "14 de septiembre de 2020", "descripcion": "Version Inicial", "asesor": "Andrea Catalina Cuesta Ruiz"}]
    cc_table = doc.add_table(rows=1, cols=4)
    cc_table.style = "Table Grid"
    cc_headers = ["Version", "Fecha", "Descripcion del cambio", "Asesor del proceso"]
    for i, h in enumerate(cc_headers):
        cc_table.rows[0].cells[i].text = h
    for item in control:
        if not isinstance(item, dict):
            item = {}
        r = cc_table.add_row().cells
        r[0].text = t(item.get("version"), "1")
        r[1].text = t(item.get("fecha"))
        r[2].text = t(item.get("descripcion"))
        r[3].text = t(item.get("asesor"))

    doc.save(docx_path)
    return docx_path


def _compilar_latex_a_pdf(tex_path):
    pdflatex = shutil.which("pdflatex")
    if not pdflatex:
        miktex_local = os.path.join(
            os.environ.get("LOCALAPPDATA", ""),
            "Programs",
            "MiKTeX",
            "miktex",
            "bin",
            "x64",
            "pdflatex.exe"
        )
        if os.path.exists(miktex_local):
            pdflatex = miktex_local
        else:
            raise RuntimeError(
                "No se encontro 'pdflatex' en el PATH. Instala MiKTeX o TeX Live para generar PDF."
            )

    workdir = os.path.dirname(tex_path) or "."
    tex_name = os.path.basename(tex_path)

    for _ in range(2):
        proc = subprocess.run(
            [pdflatex, "-interaction=nonstopmode", "-halt-on-error", tex_name],
            cwd=workdir,
            capture_output=True,
            text=True,
            timeout=180,
            check=False
        )
        if proc.returncode != 0:
            detalle = (proc.stdout or "") + "\n" + (proc.stderr or "")
            detalle = detalle[-2000:]
            raise RuntimeError(f"Fallo compilando LaTeX. Detalle:\n{detalle}")

    pdf_path = os.path.splitext(tex_path)[0] + ".pdf"
    if not os.path.exists(pdf_path):
        raise RuntimeError("La compilacion termino sin crear el PDF esperado.")

    return pdf_path


def procesar_generar_pdf_acta():
    global ultima_transcripcion

    if not ultima_transcripcion or not ultima_transcripcion.strip():
        messagebox.showerror("Error", "Primero debes transcribir un audio para generar el PDF del acta.")
        return

    text_box.insert(tk.END, "\nGenerando acta en PDF...\n")
    text_box.update()

    try:
        tex_path, uso_local_por_cuota, datos = _generar_y_guardar_acta(usar_google=True)
        text_box.insert(tk.END, f"Acta LaTeX generada: {tex_path}\n")
        text_box.update()

        docx_path = os.path.splitext(tex_path)[0] + ".docx"
        try:
            _guardar_acta_docx(datos, docx_path)
            text_box.insert(tk.END, f"Acta Word generada: {docx_path}\n")
            text_box.update()
        except Exception as e_docx:
            text_box.insert(tk.END, f"No se pudo generar Word: {e_docx}\n")
            text_box.update()

        pdf_path = _compilar_latex_a_pdf(tex_path)
        text_box.insert(tk.END, f"Acta PDF generada: {pdf_path}\n")
        text_box.update()

        if uso_local_por_cuota:
            messagebox.showwarning(
                "PDF generado con resumen local",
                "Google API no tuvo cuota disponible. Se genero el PDF usando datos locales en:\n"
                f"{pdf_path}"
            )
        else:
            messagebox.showinfo("Exito", f"Acta PDF generada en:\n{pdf_path}\n\nActa Word generada en:\n{docx_path}")
    except Exception as e:
        err_txt = str(e)
        if "pdflatex" in err_txt.lower():
            messagebox.showwarning(
                "Falta compilador LaTeX",
                f"{e}\n\nLa app dejo listo el archivo .tex para compilar manualmente."
            )
        else:
            messagebox.showerror("Error", f"No se pudo generar el PDF del acta: {e}")


def procesar_mejora():
    if not audio_file:
        messagebox.showerror("Error", "Debe seleccionar un archivo de audio.")
        return

    temp_audio = None
    try:
        audio_para_procesar, temp_created = preparar_audio_wav(audio_file, text_box)
        if temp_created:
            temp_audio = audio_para_procesar
        mejorado = mejorar_audio(audio_para_procesar)
        text_box.insert(tk.END, f"Audio mejorado generado: {mejorado}\n\n")
    except RuntimeError as e:
        text_box.insert(tk.END, f"{e}\n")
    except Exception as e:
        text_box.insert(tk.END, f"Error mejorando audio: {e}\n")
    finally:
        if temp_audio and os.path.exists(temp_audio):
            try:
                os.remove(temp_audio)
            except Exception:
                pass


def procesar_calidad():
    if not audio_file:
        messagebox.showerror("Error", "Debe seleccionar un archivo de audio.")
        return

    temp_audio = None
    try:
        audio_para_procesar, temp_created = preparar_audio_wav(audio_file, text_box)
        if temp_created:
            temp_audio = audio_para_procesar
        calidad = verificar_calidad(audio_para_procesar)
        text_box.insert(tk.END, f"Calidad del audio: {calidad}\n\n")
    except RuntimeError as e:
        text_box.insert(tk.END, f"{e}\n")
    finally:
        if temp_audio and os.path.exists(temp_audio):
            try:
                os.remove(temp_audio)
            except Exception:
                pass


def procesar_voces():
    if not audio_file:
        messagebox.showerror("Error", "Debe seleccionar un archivo de audio.")
        return

    temp_audio = None
    try:
        audio_para_procesar, temp_created = preparar_audio_wav(audio_file, text_box)
        if temp_created:
            temp_audio = audio_para_procesar
        resultado = identificar_voces(audio_para_procesar)
        text_box.insert(tk.END, f"Voces identificadas: {resultado}\n\n")
    except RuntimeError as e:
        text_box.insert(tk.END, f"{e}\n")
    finally:
        if temp_audio and os.path.exists(temp_audio):
            try:
                os.remove(temp_audio)
            except Exception:
                pass


# ------------------------------------
# Ventana principal
# ------------------------------------

root = tk.Tk()
root.title("Transcriptor GPLv3 - Español")
root.geometry("800x600")

audio_file = None
ultima_transcripcion = ""
api_key_var = tk.StringVar(value=os.environ.get("GEMINI_API_KEY", ""))

btn_select = tk.Button(root, text="Seleccionar Audio", command=seleccionar_audio, width=30)
btn_select.pack(pady=10)

btn_quality = tk.Button(root, text="Verificar Calidad", command=procesar_calidad, width=30)
btn_quality.pack(pady=10)

btn_improve = tk.Button(root, text="Mejorar Audio (Eliminar Ruido)", command=procesar_mejora, width=30)
btn_improve.pack(pady=10)

btn_speakers = tk.Button(root, text="Identificar Voces", command=procesar_voces, width=30)
btn_speakers.pack(pady=10)

btn_transcribe = tk.Button(root, text="Transcribir Audio", command=procesar_transcripcion, width=30)
btn_transcribe.pack(pady=10)

api_key_label = tk.Label(root, text="Google API Key (Gemini):")
api_key_label.pack(pady=(5, 0))

api_key_entry = tk.Entry(root, textvariable=api_key_var, width=60, show="*")
api_key_entry.pack(pady=5)

btn_summary = tk.Button(root, text="Resumir Reunion (Google AI)", command=procesar_resumen_reunion, width=30)
btn_summary.pack(pady=10)

btn_acta = tk.Button(root, text="Generar Acta LaTeX", command=procesar_generar_acta_latex, width=30)
btn_acta.pack(pady=10)

btn_acta_pdf = tk.Button(root, text="Generar Acta PDF + Word", command=procesar_generar_pdf_acta, width=30)
btn_acta_pdf.pack(pady=10)

text_box = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=90, height=20)
text_box.pack(padx=10, pady=10)

root.mainloop()
